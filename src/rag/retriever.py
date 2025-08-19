import glob
import os
import io
import sys
from typing import List, Optional
import tempfile

from langchain_community.retrievers import BM25Retriever
from langchain_core.documents import Document
from langchain_core.retrievers import BaseRetriever
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from pydantic import Field, BaseModel
from llm.config import get_gemini_llm

# Optional imports for file processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Set UTF-8 encoding
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stdin = io.TextIOWrapper(sys.stdin.buffer, encoding='utf-8')


class HybridRetriever(BaseRetriever, BaseModel):
    vectorstore: FAISS = Field(description="FAISS vector store")
    bm25_retriever: BM25Retriever = Field(description="BM25 retriever")
    k: int = Field(default=4, description="Number of documents to retrieve")

    class Config:
        arbitrary_types_allowed = True

    def _get_relevant_documents(self, query: str) -> List[Document]:
        vector_docs = self.vectorstore.similarity_search(query, k=self.k)
        bm25_docs = self.bm25_retriever.get_relevant_documents(query)

        all_docs = []
        seen_content = set()
        for doc in vector_docs + bm25_docs:
            if doc.page_content not in seen_content:
                all_docs.append(doc.page_content)
                seen_content.add(doc.page_content)

        return [Document(page_content=content) for content in all_docs]

    async def _aget_relevant_documents(self, query: str) -> List[Document]:
        raise NotImplementedError("Async retrieval not implemented")


def read_all_text_files(data_dir):
    """Đọc toàn bộ nội dung các file .txt trong thư mục"""
    combined_text = ""
    for file_path in glob.glob(os.path.join(data_dir, "*.txt")):
        with open(file_path, "r", encoding="utf-8") as f:
            combined_text += f.read() + "\n\n"
    return combined_text


def create_vector_database(output_path, data_dir="./data"):
    try:
        regulations = read_all_text_files(data_dir)

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=400, 
            chunk_overlap=200, 
            length_function=len,
            separators=["\n\n", "\n", " ", ""],  
            keep_separator=False
        )
        chunks = text_splitter.split_text(regulations)

        embeddings = OllamaEmbeddings(
            model="nomic-embed-text",
            base_url="http://42.112.213.93:11434"
        )
        # embeddings = OllamaEmbeddings(
        #     model="nomic-embed-text"
        # )

        vectorstore = FAISS.from_texts(chunks, embeddings)

        if os.path.dirname(output_path):
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
        else:
            os.makedirs(output_path, exist_ok=True)

        vectorstore.save_local(output_path)
        return chunks
    except Exception as e:
        print(f"Error creating vector database: {e}")
        raise


def load_vector_database(output_path, data_dir="./data"):
    try:
        embeddings = OllamaEmbeddings(
            model="nomic-embed-text",
            base_url="http://42.112.213.93:11434"
        )
        # embeddings = OllamaEmbeddings(
        #     model="nomic-embed-text"
        # )

        if not os.path.exists(output_path):
            chunks = create_vector_database(output_path, data_dir)
        else:
            regulations = read_all_text_files(data_dir)
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=400, 
                chunk_overlap=200, 
                length_function=len,
                separators=["\n\n", "\n", " ", ""], 
                keep_separator=False
            )
            chunks = text_splitter.split_text(regulations)

        print("Loading vector database...")
        return FAISS.load_local(output_path, embeddings, allow_dangerous_deserialization=True), chunks
    except Exception as e:
        print(f"Error loading vector database: {e}")
        raise


def create_hybrid_retriever(vector_db_path, data_dir="./data"):
    vectorstore, documents = load_vector_database(vector_db_path, data_dir)
    bm25_retriever = BM25Retriever.from_texts(texts=documents, k=15)

    return HybridRetriever(
        vectorstore=vectorstore, 
        bm25_retriever=bm25_retriever, 
        k=15
    ), documents


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text from uploaded file based on file type"""
    try:
        if file_type == "text/plain" or file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_type == "application/pdf" or file_path.endswith('.pdf'):
            # Try to import PyPDF2 for PDF processing
            try:
                import PyPDF2
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text
            except ImportError:
                return "PDF processing not available. Please install PyPDF2."
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_path.endswith('.docx'):
            # Try to import python-docx for Word document processing
            try:
                import docx
                doc = docx.Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                return "DOCX processing not available. Please install python-docx."
        
        else:
            return "Unsupported file type. Please upload .txt, .pdf, or .docx files."
    
    except Exception as e:
        return f"Error extracting text from file: {str(e)}"


def create_in_memory_retriever(file_content: str, chunk_size: int = 400, chunk_overlap: int = 200, k: int = 15) -> HybridRetriever:
    """Create an in-memory hybrid retriever from file content"""
    try:
        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""],
            keep_separator=False
        )
        chunks = text_splitter.split_text(file_content)
        
        # Create embeddings
        embeddings = OllamaEmbeddings(
            model="nomic-embed-text"
        )
        
        # Create in-memory FAISS vector store
        vectorstore = FAISS.from_texts(chunks, embeddings)
        
        # Create BM25 retriever
        bm25_retriever = BM25Retriever.from_texts(texts=chunks, k=k)
        
        # Create hybrid retriever
        hybrid_retriever = HybridRetriever(
            vectorstore=vectorstore,
            bm25_retriever=bm25_retriever,
            k=k
        )
        
        return hybrid_retriever, chunks
    
    except Exception as e:
        print(f"Error creating in-memory retriever: {e}")
        raise


def process_uploaded_file(uploaded_file, chunk_size: int = 400, chunk_overlap: int = 200, k: int = 15):
    """Process an uploaded file and create an in-memory retriever"""
    try:
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        # Extract text from file
        file_content = extract_text_from_file(tmp_file_path, uploaded_file.type)
        
        # Clean up temporary file
        os.unlink(tmp_file_path)
        
        if file_content.startswith("Error") or file_content.startswith("Unsupported") or "not available" in file_content:
            return None, file_content
        
        # Create in-memory retriever
        retriever, chunks = create_in_memory_retriever(file_content, chunk_size, chunk_overlap, k)
        
        return retriever, f"Successfully processed file: {uploaded_file.name}. Created {len(chunks)} chunks."
    
    except Exception as e:
        return None, f"Error processing file: {str(e)}"
